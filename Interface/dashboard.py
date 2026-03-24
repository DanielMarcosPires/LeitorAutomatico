# --- BIBLIOTECAS PADRÃO (Nativas do Python) ---
import os
import sys
import glob
from datetime import datetime

# --- BIBLIOTECAS EXTERNAS (Instaladas via pip) ---
import pandas as pd
import openpyxl
import customtkinter
import tksvg
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

# --- MÓDULOS INTERNOS (DaniTechnologia) ---
from Interface.Fonts.fonts import Fonts
from Interface.colors import DashboardColors
from Interface.Classes.folders import folder

class Dashboard(customtkinter.CTkFrame):
    fonts = Fonts()
    
    def __init__(self, master, **kwargs):
        super().__init__(master, **kwargs)
        self.master.after(0, lambda: self.master.state("zoomed"))  # type: ignore 
          
        # Carregar ícones
        self._load_icons()
        
        # Configurar o título principal
        self.title()
        
        # Configurar o layout principal com frames esquerdo e direito
        self.main_layout()
        
        # Configurar os botões no frame esquerdo
        self.buttons()
        
        # Configurar a lista de planilhas no frame direito
        self.excel_list()
        
    # ==========================================
    # 1. MÉTODOS DE INTERFACE (UI/UX)
    # ==========================================
    # Tudo que desenha botões, frames e o foguete
    
    def title(self):
        """Configura o título principal com botão Sair na direita."""
        title_frame = customtkinter.CTkFrame(
            self,
            height=50,
            fg_color=DashboardColors.TITLE_BACKGROUND,
            border_width=0
        )
        title_frame.pack(fill="x", padx=20, pady=(20, 20))
        title_frame.pack_propagate(False)

        # Título à esquerda
        ctkTitle = customtkinter.CTkLabel(
            title_frame,
            text="Dashboard - Gerenciar Planilhas",
            font=self.fonts.MAIN_TITLE,
            text_color=DashboardColors.TITLE_TEXT
        )
        ctkTitle.pack(side="left", padx=20, pady=5)

        # Botão Sair à direita
        btn_sair = customtkinter.CTkButton(
            title_frame,
            text="Sair",
            image=self.icons.get('exit'),
            fg_color=DashboardColors.BUTTON_CRITICAL_BACKGROUND,
            hover_color="#c41e3a",
            text_color=DashboardColors.BUTTON_TEXT,
            width=80,
            height=35,
            font=self.fonts.BUTTON_EXIT,
            compound="left",
            command=self.logout
        )
        btn_sair.pack(side="right", padx=20, pady=5)
    
    def logoText(self,text:str,fontFamily:str=None,weight:str=None): # pyright: ignore[reportArgumentType]
        ctkTitle = customtkinter.CTkTextbox(
            self.left_frame, 
            height=30, 
            font=self.fonts.LOGO if not fontFamily else (fontFamily, 24, weight or "bold"), # type: ignore
            text_color=DashboardColors.LEFT_TITLE_TEXT
        )
        ctkTitle.tag_config("center", justify='center')
        ctkTitle.insert("0.0", text)
        ctkTitle.tag_add("center", "1.0", "end")
        ctkTitle.configure(state="disabled")
        ctkTitle.pack(fill="x", padx=20, pady=(20, 10))
    
    def buttons(self):
        """Configura os botões de ação no painel esquerdo."""
        # Botão Gerar Relatório
        
        btn_relatorio = customtkinter.CTkButton(
            self.left_frame,
            text="Gerar Relatório",
            image=self.icons.get('relatory'),
            fg_color=DashboardColors.BUTTON_CRITICAL_BACKGROUND,
            hover_color=DashboardColors.BUTTON_HOVER,
            text_color=DashboardColors.BUTTON_TEXT,
            width=170,
            height=40,
            font=self.fonts.BUTTON_LARGE,
            compound="left",
            anchor="w",
            command=self.create_report
        )
        btn_relatorio.pack(side="bottom",fill="x", pady=6, anchor="w", padx=8)


        # Botão Visualizar
        btn_visualizar = customtkinter.CTkButton(
            self.left_frame,
            text="Visualizar",
            image=self.icons.get('view'),
            fg_color=DashboardColors.BUTTON_BACKGROUND_SECONDARY,
            hover_color=DashboardColors.BUTTON_HOVER,
            text_color=DashboardColors.BUTTON_TEXT,
            width=170,
            height=40,
            font=self.fonts.BUTTON_LARGE,
            compound="left",
            anchor="w",
            command=self.view_excel_content
        )
        btn_visualizar.pack(side="bottom",fill="x", pady=6, anchor="w", padx=8)

        # Botão Atualizar
        btn_atualizar = customtkinter.CTkButton(
            self.left_frame,
            text="Atualizar",
            image=self.icons.get('reload'),
            fg_color=DashboardColors.BUTTON_BACKGROUND_SECONDARY,
            hover_color=DashboardColors.BUTTON_HOVER,
            text_color=DashboardColors.BUTTON_TEXT,
            width=170,
            height=40,
            font=self.fonts.BUTTON_LARGE,
            compound="left",
            anchor="w",
            command=self.load_excel_files
        )
        btn_atualizar.pack(side="bottom",fill="x", pady=6, anchor="w", padx=8)
    
    def left_title(self):
        """Configura o título do painel esquerdo com estilo Excel."""
        ctkTitle = customtkinter.CTkTextbox(
            self.left_frame, 
            height=30, 
            font=self.fonts.SECTION_TITLE, 
            text_color=DashboardColors.LEFT_TITLE_TEXT
        ) 
        ctkTitle.tag_config("center", justify='center')
        ctkTitle.insert("0.0", "Ações")
        ctkTitle.tag_add("center", "1.0", "end")
        ctkTitle.configure(state="disabled")
        ctkTitle.pack(fill="x", padx=20, pady=(20, 10))
    
    def adesivos(self):
        # Caminho do arquivo
        svg_path = os.path.join(os.path.dirname(__file__), 'icon', 'aceleraRocket.svg')
        
        # Criamos a imagem SVG compatível com Tkinter
        # Nota: O CTkLabel aceita imagens do PhotoImage padrão do tksvg
        self.adesivo_image = tksvg.SvgImage(file=svg_path, scale=1.0) # Ajuste o scale se necessário

        self.adesivo_label = customtkinter.CTkLabel(
            self.main_frame,
            image=self.adesivo_image, # type: ignore
            text="",
            fg_color="transparent" # Garante que não haja fundo
        )
        # Para garantir que ele fique NO TOPO de tudo:
        self.adesivo_label.place(x=70, y=150)
    
    def main_layout(self):
        """Configura o layout principal com cores do Excel."""
        self.main_frame = customtkinter.CTkFrame(
            self,
            fg_color=DashboardColors.MAIN_FRAME_BACKGROUND,  # Fundo escuro como Excel
            border_width=1,
            border_color=DashboardColors.MAIN_FRAME_BORDER  # Borda azul Excel
        )
        self.main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Frame esquerdo para botões
        self.left_frame = customtkinter.CTkFrame(
            self.main_frame, 
            width=300,
            fg_color=DashboardColors.LEFT_FRAME_BACKGROUND,  # Fundo escuro
            border_width=1,
            border_color=DashboardColors.LEFT_FRAME_BORDER  # Borda azul Excel
        )
        self.left_frame.pack(side="left", fill="y", padx=10, pady=10)
        self.left_frame.pack_propagate(False)  # Manter largura fixa

        self.logoDevice()  # Logo personalizada

        self.adesivos()  # Adesivo voando no meio da tela

        # Frame direito para lista de planilhas
        self.right_frame = customtkinter.CTkFrame(
            self.main_frame,
            fg_color=DashboardColors.RIGHT_FRAME_BACKGROUND,  # Fundo escuro
            border_width=1,
            border_color=DashboardColors.RIGHT_FRAME_BORDER  # Borda azul Excel
        )
        self.right_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)
    
    def logoDevice(self):
        """Unifica ACELERA e G&P na mesma linha corrigindo o erro de fonte"""
        ctkTitle = customtkinter.CTkTextbox(
            self.left_frame, 
            fg_color="transparent",
            height=50,
            text_color=DashboardColors.LEFT_TITLE_TEXT,
            activate_scrollbars=False,
            wrap="none"
        )
        
        # Criamos as fontes (o CTkFont ajuda a manter o scaling no restante do app)
        font_acelera = customtkinter.CTkFont(family="Montserrat", size=28, weight="bold")
        font_gp = customtkinter.CTkFont(family="Inter", size=24, weight="bold")

        # ACESSO DIRETO AO TKINTER:
        # Usamos ctkTitle._textbox para configurar as tags diretamente no widget base
        ctkTitle._textbox.tag_config("font_acelera", font=font_acelera, foreground="#f9c319")
        ctkTitle._textbox.tag_config("font_gp", font=font_gp, foreground="#fff")
        ctkTitle._textbox.tag_config("center", justify='center')

        # Inserção
        ctkTitle.insert("0.0", "ACELERA", "font_acelera")
        ctkTitle.insert("end", " G&P", "font_gp") 
        
        # Aplica a centralização em tudo
        ctkTitle.tag_add("center", "1.0", "end")
        
        ctkTitle.configure(state="disabled")
        ctkTitle.pack(fill="x", padx=20, pady=(20, 10))  
    
    def excel_list(self):
        """Configura a lista de planilhas no painel direito com título e botões no topo."""
        # Frame para título e botões na mesma linha
        header_frame = customtkinter.CTkFrame(
            self.right_frame,
            fg_color="transparent",
            border_color=DashboardColors.RIGHT_FRAME_BORDER,
            border_width=1,
            height=50,
        )
        header_frame.pack(fill="x")
        header_frame.pack_propagate(False)

        # Título à esquerda
        self.right_title = customtkinter.CTkLabel(
            header_frame,
            text="Planilhas Disponíveis",
            font=self.fonts.SECTION_TITLE,
            text_color=DashboardColors.RIGHT_TITLE_TEXT
        )
        self.right_title.pack(side="left", padx=20, pady=0, anchor="center")

        # Frame de botões à direita
        buttons_frame = customtkinter.CTkFrame(
            header_frame,
            fg_color="transparent",
            border_width=0
        )
        buttons_frame.pack(side="right", padx=20, pady=0)

        # Botão Criar nova planilha
        btn_criar = customtkinter.CTkButton(
            buttons_frame,
            text="",
            image=self.icons.get('excel_logo'),
            fg_color=DashboardColors.BUTTON_BACKGROUND_SECONDARY,
            hover_color=DashboardColors.BUTTON_HOVER,
            text_color=DashboardColors.BUTTON_TEXT,
            width=90,
            height=30,
            font=self.fonts.BUTTON_SMALL,
            compound="left",
            command=self.create_new_excel
        )
        btn_criar.pack(side="left", padx=5)

        # Botão Excluir
        btn_excluir = customtkinter.CTkButton(
            buttons_frame,
            text="",
            image=self.icons.get('trash'),
            fg_color=DashboardColors.BUTTON_BACKGROUND_SECONDARY,
            hover_color=DashboardColors.BUTTON_HOVER,
            text_color=DashboardColors.BUTTON_TEXT,
            width=90,
            height=30,
            font=self.fonts.BUTTON_SMALL,
            command=self.delete_excel_files
        )
        btn_excluir.pack(side="left", padx=5)

        # Botão Abrir pasta
        btn_abrir = customtkinter.CTkButton(
            buttons_frame,
            text="",
            image=self.icons.get('folder'),
            fg_color=DashboardColors.BUTTON_BACKGROUND_SECONDARY,
            hover_color=DashboardColors.BUTTON_HOVER,
            text_color=DashboardColors.BUTTON_TEXT,
            width=90,
            height=30,
            font=self.fonts.BUTTON_SMALL,
            command=self.open_folder_planilhas
        )
        btn_abrir.pack(side="left", padx=5)

        # Scrollable frame para a lista
        self.scrollable_frame = customtkinter.CTkScrollableFrame(
            self.right_frame,
            height=300,
            fg_color=DashboardColors.SCROLL_BACKGROUND,
            border_width=1,
            border_color=DashboardColors.SCROLL_BORDER
        )
        self.scrollable_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Dicionário para armazenar checkboxes e seus caminhos
        self.checkboxes = {}
        self.checkbox_frames = {}

        # Carregar planilhas
        self.load_excel_files()

        # Área para resultados
        self.result_textbox = customtkinter.CTkTextbox(
            self.right_frame,
            height=200,
            fg_color=DashboardColors.RESULT_TEXTBOX_BACKGROUND,
            border_width=1,
            border_color=DashboardColors.RESULT_TEXTBOX_BORDER,
            text_color=DashboardColors.RESULT_TEXTBOX_TEXT
        )
        self.result_textbox.pack(fill="x", padx=10, pady=10)
    
    def excel_list_hoverItem(self, frame, is_hover):
        """
        Aplica efeito de hover nos itens da lista de planilhas usando cores do Excel:
        """
        if is_hover:
            # Salvar cor atual antes do hover
            current_color = frame.cget("fg_color")
            frame._original_color = current_color
            frame.configure(fg_color="#383700", border_color="#FFD133")  # Azul Excel no hover
        else:
            # Restaurar cor original
            if hasattr(frame, '_original_color'):
                frame.configure(fg_color="#383700", border_color="#ffd133")
    
    def excelItem_checkbox_activeColor(self, frame, checkbox):
        """
        Atualiza a cor de fundo do frame baseado no estado do checkbox.
        
        Usa paletas de cores do Excel no tema escuro:
        - Não selecionado: Fundo escuro (#2a2a2a)
        - Selecionado: Azul Excel (#4F81BD)
        - Hover: Azul mais claro (#6B9BD1)
        """
        if checkbox.get() == 1:  # Selecionado
            frame.configure(fg_color="#383700", border_color="#ffd133")  # Azul Excel para selecionado
        else:  # Não selecionado
            frame.configure(fg_color="#2a2a2a", border_color="#404040")  # Fundo escuro
    
    # =========================================
    # 2. MÉTODOS DE AÇÃO (LÓGICA DE NEGÓCIO)
    # =========================================
    # Tudo que processa dados, lê arquivos, gera relatórios, etc.
    
    def _load_icons(self):
        """Carrega os ícones dos botões a partir da pasta Interface/icon/."""
        self.icons = {}
        
        # Mapeamento de nomes de ícones para arquivos
        icon_files = {
            'relatory': 'relatory.png',
            'view': 'view.png', 
            'reload': 'reload.png',
            'exit': 'exit.png',
            'excel_logo': 'cells.png',
            'trash': 'trash.png',
            'folder': 'folder.png'
        }
        
        # Caminho base para os ícones
        icon_base_path = os.path.join(os.path.dirname(__file__), 'icon')
        
        for icon_name, filename in icon_files.items():
            icon_path = os.path.join(icon_base_path, filename)
            try:
                # Carregar imagem com PIL e redimensionar para 20x20 pixels
                pil_image = Image.open(icon_path)
                pil_image = pil_image.resize((20, 20), Image.Resampling.LANCZOS)
                self.icons[icon_name] = customtkinter.CTkImage(light_image=pil_image, dark_image=pil_image, size=(20, 20))
            except (FileNotFoundError, OSError) as e:
                print(f"Aviso: Ícone '{filename}' não encontrado em {icon_path}. Usando texto apenas.")
                self.icons[icon_name] = None
    
    def open_folder_planilhas(self):
        """Abre a pasta Planilhas/ no explorador de arquivos."""
        planilhas_dir = os.path.join(os.getcwd(), "Planilhas")
        if not os.path.exists(planilhas_dir):
            os.makedirs(planilhas_dir)
        os.startfile(planilhas_dir)  # Abrir pasta no explorador de arquivos (Windows)    
    
    
    def logout(self):
        """Retorna à tela de login."""
        self.pack_forget()
        self.master.meu_form.pack(pady=20, expand=True)  # type: ignore
    
    def load_excel_files(self):
        """
        Carrega e exibe a lista de arquivos Excel (.xlsx) encontrado na pasta Planilhas.
        
        Procura recursivamente por arquivos .xlsx no diretório atual
        e subdiretórios, criando frames com checkboxes para seleção.
        Cada item tem background que muda quando selecionado.
        """
        # Limpar checkboxes anteriores
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
        self.checkboxes = {}
        self.checkbox_frames = {}  # Novo dicionário para armazenar os frames
        
        # Procurar arquivos .xlsx no diretório atual e subdiretórios
        excel_files = glob.glob("Planilhas/*.xlsx", recursive=True)
        
        if not excel_files:
            no_files_label = customtkinter.CTkLabel(self.scrollable_frame, text="Nenhuma planilha encontrada.")
            no_files_label.pack(pady=20)
            return
        
        for file_path in excel_files:
            # Criar frame para cada item com background Excel
            item_frame = customtkinter.CTkFrame(
                self.scrollable_frame,
                fg_color="#2a2a2a",  # Fundo escuro Excel (#2a2a2a)
                corner_radius=5,
                border_width=1,
                border_color="#404040"  # Borda sutil
            )
            item_frame.pack(fill="x", padx=5, pady=2)
            
            # Adicionar eventos de mouse para hover effect
            item_frame.bind("<Enter>", lambda e, f=item_frame: self.excel_list_hoverItem(f, True))
            item_frame.bind("<Leave>", lambda e, f=item_frame: self.excel_list_hoverItem(f, False))
            
            # Checkbox dentro do frame
            checkbox = customtkinter.CTkCheckBox(
                item_frame, 
                text=os.path.basename(file_path)
            )
            checkbox.pack(anchor="w", padx=10, pady=5)
            
            # Configurar callback para mudança de background
            checkbox.configure(command=lambda f=item_frame, c=checkbox: self.excelItem_checkbox_activeColor(f, c))
            
            # Armazenar referências
            self.checkboxes[checkbox] = file_path
            self.checkbox_frames[checkbox] = item_frame
    
    def create_new_excel(self):
        """
        Criar uma nova planilha de alunos em Excel na pasta Planilhas.
        """
        planilhas_dir = os.path.join(os.getcwd(), "Planilhas")
        if not os.path.exists(planilhas_dir):
            os.makedirs(planilhas_dir)
        
        new_file_path = os.path.join(planilhas_dir, "nova_planilha.xlsx")
        if os.path.exists(new_file_path):
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", "Arquivo nova_planilha.xlsx já existe.")
            return
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Alunos" # type: ignore
        
        # Configurar cabeçalho com estilo Excel
        headers = ["NOME",	"ETAPA",	"Critério",	"STATUS",	"Qualidade",	"Comunicação",	"Aprendizado",	"Conhecimento"]
        for col_num, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_num, value=header) # type: ignore
            cell.font = Font(bold=True, color="FFFFFF")  # Texto branco
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")  # Azul Excel
        
        wb.save(new_file_path)
        
        self.result_textbox.delete("0.0", "end")
        self.result_textbox.insert("0.0", f"Nova planilha criada: {new_file_path}")
        
        # Atualizar lista de planilhas
        self.load_excel_files()
    
    def delete_excel_files(self):
        """
        Exclui os arquivos Excel selecionados via checkboxes e é transferido para pasta DeleteFiles.
        """
        selected_files = [path for checkbox, path in self.checkboxes.items() if checkbox.get() == 1]
        
        if not selected_files:
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", "Nenhuma planilha selecionada para exclusão.")
            return
        
        delete_dir = os.path.join(os.getcwd(), "DeleteFiles")
        if not os.path.exists(delete_dir):
            os.makedirs(delete_dir)
        
        for file_path in selected_files:
            try:
                # Mover arquivo para pasta DeleteFiles
                os.rename(file_path, os.path.join(delete_dir, os.path.basename(file_path)))
                self.result_textbox.insert("end", f"Arquivo movido para DeleteFiles: {os.path.basename(file_path)}\n")
            except Exception as e:
                self.result_textbox.insert("end", f"Erro ao mover {os.path.basename(file_path)}: {str(e)}\n")
        
        # Atualizar lista de planilhas
        self.load_excel_files()
    
    def read_selected_excels(self):
        """
        Lê os arquivos Excel selecionados e conta os nomes de clientes.
        
        Processa múltiplas planilhas selecionadas via checkboxes,
        extraindo nomes da segunda coluna e exibindo contagem.
        """
        selected_files = [path for checkbox, path in self.checkboxes.items() if checkbox.get() == 1]
        
        if not selected_files:
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", "Nenhuma planilha selecionada.")
            return
        
        results = []
        for file_path in selected_files:
            try:
                wb = openpyxl.load_workbook(file_path)
                sheet = wb.active
                names = [row[0] for row in sheet.iter_rows(min_row=2, values_only=True) if row[0]] # type: ignore
                results.append(f"{os.path.basename(file_path)}: {len(names)} nomes lidos")
            except Exception as e:
                results.append(f"Erro ao ler {os.path.basename(file_path)}: {str(e)}")
        
        self.result_textbox.delete("0.0", "end")
        self.result_textbox.insert("0.0", "\n".join(results))
    
    def generate_folders(self):
        """Gera pastas com o nome da tabela e subpastas para cada nome listado na primeira coluna das planilhas selecionadas."""
        selected_files = [path for checkbox, path in self.checkboxes.items() if checkbox.get() == 1]
        
        if not selected_files:
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", "Nenhuma planilha selecionada para gerar pastas.")
            return
        
        
        for file_path in selected_files:
            try:
                wb = openpyxl.load_workbook(file_path)
                sheet = wb.active
                names = [row[0] for row in sheet.iter_rows(min_row=2, values_only=True) if row[0]] # type: ignore
                if names:
                    folder_handler = folder(sheet.title) # pyright: ignore[reportOptionalMemberAccess]
                    folder_handler.generateFolders(sheet.title, names) # type: ignore
                    self.result_textbox.insert("end", f"Pastas geradas para {os.path.basename(file_path)}\n")
                else:
                    self.result_textbox.insert("end", f"Planilha {os.path.basename(file_path)} vazia.\n")
            except Exception as e:
                self.result_textbox.insert("end", f"Erro em {os.path.basename(file_path)}: {str(e)}\n")
    
    def create_report(self):
        """
        Cria a estrutura: Projetos > [NomeDaPlanilha] > [NomeDaPessoa] > Relatorio.docx
        """
        selected_files = [path for checkbox, path in self.checkboxes.items() if checkbox.get() == 1]
        
        if not selected_files:
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", "Nenhuma planilha selecionada.")
            return

        summary = []
        
        for file_path in selected_files:
            try:
                project_name = os.path.splitext(os.path.basename(file_path))[0]
                wb = openpyxl.load_workbook(file_path, data_only=True)
                sheet = wb.active
                
                count_processed = 0

                for row in sheet.iter_rows(min_row=2, values_only=True): # type: ignore
                    if not row or row[0] is None:
                        continue
                    
                    nome_pessoa = str(row[0]).strip()
                    
                    # Coleta de dados baseada na sua imagem
                    dados = {
                        "Etapa": str(row[1]) if len(row) > 1 else "-",
                        "Status": str(row[3]) if len(row) > 3 else "-",
                        "Qualidade": str(row[4]) if len(row) > 4 else "-",
                        "Comunicação": str(row[5]) if len(row) > 5 else "-",
                        "Aprendizado": str(row[6]) if len(row) > 6 else "-",
                        "Conhecimento": str(row[7]) if len(row) > 7 else "-"
                    }
                    # --- LÓGICA DO TEXTO DINÂMICO ---
                    if "Férias" in dados["Status"] or "Férias" in dados["Etapa"]:
                        desc_texto = "Registro de pausa programada. Devido ao período de férias, não há dados de performance registrados. A tabela reflete apenas o status administrativo."
                    elif "Não entregue" in dados["Etapa"] or "atraso" in dados["Status"].lower():
                        desc_texto = "Este relatório indica que a atividade ainda não atingiu o estado de conclusão esperado. É fundamental a regularização dos pontos abaixo para nova avaliação."
                    elif "Ótimo" in dados["Qualidade"] and "Ótimo" in dados["Comunicação"]:
                        desc_texto = "O colaborador apresenta um aproveitamento integral das competências. Os indicadores abaixo demonstram alto nível de clareza e domínio técnico."
                    else:
                        desc_texto = "A etapa foi concluída e os requisitos atendidos. A tabela abaixo detalha o equilíbrio entre os critérios técnicos e oportunidades de refinamento."                    

                    # Criar diretórios
                    person_dir = os.path.join('Projetos', project_name, nome_pessoa)
                    os.makedirs(person_dir, exist_ok=True)

                    # --- GERAÇÃO DO ARQUIVO WORD ---
                    doc = Document()

                    # Configuração de Fonte Global (Arial para profissionalismo)
                    style = doc.styles['Normal']
                    font = style.font # type: ignore
                    font.name = 'Roboto'
                    font.size = Pt(11)

                    # Cabeçalho de Identificação
                    p_id = doc.add_paragraph()
                    run_relatoryTitle = p_id.add_run("Relatório de desempenho:\n")
                    run_relatoryTitle.font.size = Pt(20)
                    run_relatoryTitle.font.bold = True
                    
                    run_colaborador = p_id.add_run("Colaborador(a): ")
                    run_colaborador.font.bold = True
                    run_colaborador_name = p_id.add_run(f"{nome_pessoa}\n")
                    
                    run_colaborador.font.size = Pt(16)
                    run_colaborador_name.font.size = Pt(16)
                    
                    run_project = p_id.add_run("Projeto: ")
                    run_project.font.bold = True
                    run_project_name = p_id.add_run(f"{project_name}\n")
                    
                    run_project.font.size = Pt(16)
                    run_project_name.font.size = Pt(16)
                    
                    run_data = p_id.add_run("Data de Emissão: ")
                    run_data.font.bold = True
                    run_data_value = p_id.add_run(f"{pd.Timestamp.now().strftime('%d/%m/%Y')} às {pd.Timestamp.now().strftime('%H:%M')}\n")
                    
                    run_data.font.size = Pt(16)
                    run_data_value.font.size = Pt(16)
                    
                    # Texto de Introdução
                    intro = doc.add_paragraph()
                    title_info = intro.add_run("Informações de desempenho:")
                    title_info.bold = True
                    title_info.font.size = Pt(20)
                    
                    intro.add_run("\n" + desc_texto).font.size = Pt(16) # O texto dinâmico baseado nos dados da planilha

                    # --- TABELA DE AVALIAÇÃO ESTILIZADA ---
                    # Usando um estilo mais limpo que o 'Table Grid'
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light Grid Accent 1' 

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'CRITÉRIO'
                    hdr_cells[1].text = 'AVALIAÇÃO'

                    # Preenchendo os dados técnicos com lógica de cores
                    for chave, valor in dados.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(chave).upper()
                        
                        # Adicionando o valor com destaque
                        p_celula = row_cells[1].paragraphs[0]
                        run_valor = p_celula.add_run(str(valor))
                        
                        # Aplicação de cores baseada no seu padrão de dados
                        if "Ótimo" in str(valor):
                            run_valor.font.color.rgb = RGBColor(0, 128, 0) # Verde
                            run_valor.bold = True
                        elif "Regular" in str(valor):
                            run_valor.font.color.rgb = RGBColor(128, 0, 128) # Roxo (seu padrão)
                        elif "Com atraso" in str(valor) or "Não" in str(valor):
                            run_valor.font.color.rgb = RGBColor(200, 0, 0) # Vermelho

                    doc.add_paragraph() # Espaço pós tabela

                    # Seção de Conclusão
                    conclusao = doc.add_paragraph()
                    textConclusao = conclusao.add_run("Observações: ")
                    textConclusao.bold = True
                    textConclusao.font.size = Pt(16)
                    textConclusao_info =conclusao.add_run("Os critérios acima são fundamentais para o acompanhamento da qualidade e evolução do projeto.")
                    textConclusao_info.font.size = Pt(16)

                    # --- RODAPÉ DE AUTORIA ---
                    doc.add_paragraph("\n") # Um espaço em branco antes do rodapé
                    # 1. Acessa a primeira seção do documento
                    section = doc.sections[0]

                    # 2. Acessa o rodapé daquela seção
                    footer = section.footer

                    # 3. O rodapé já vem com um parágrafo padrão, vamos usá-lo ou adicionar um novo
                    # Se quiser limpar o que já existe:
                    p_footer = footer.paragraphs[0] 
                    p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # 4. Adiciona o seu texto de autoria
                    run_msg = p_footer.add_run("Relatório gerado pelo LeitorDePlanilhas desenvolvido pelo autor: \n")
                    run_msg.font.size = Pt(10)
                    run_msg.italic = True

                    run_autor = p_footer.add_run("DaniTechnologia/Daniel Marcos Pires")
                    run_autor.font.size = Pt(11)
                    run_autor.bold = True
                    run_autor.font.color.rgb = RGBColor(31, 73, 125) # Azul escuro
                    # Salvar o arquivo .docx
                    doc_path = os.path.join(person_dir, f'Relatorio_{nome_pessoa}.docx')
                    doc.save(doc_path)
                    # -------------------------------

                    count_processed += 1

                summary.append(f"Sucesso: {project_name} ({count_processed} Words gerados)")

            except Exception as e:
                summary.append(f"Erro em {os.path.basename(file_path)}: {str(e)}")

        self.result_textbox.delete("0.0", "end")
        self.result_textbox.insert("0.0", "\n".join(summary))

    def view_excel_content(self):
        """
            Visualiza o conteúdo da planilha Excel selecionada em uma nova janela com estilo inspirado no Excel.
        """
        # 1. PEGAR OS ARQUIVOS SELECIONADOS (Correção do erro de variável)
        selected_files = [path for checkbox, path in self.checkboxes.items() if checkbox.get() == 1]
        
        if not selected_files:
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", "Nenhuma planilha selecionada para visualizar.")
            return
        
        file_path = selected_files[0]  # Pega a primeira selecionada
        
        try:
            # Carrega a planilha
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheet = wb.active
            
            # Criar janela de visualização com tema Excel
            view_window = customtkinter.CTkToplevel(self)
            view_window.title(f"DaniTechnologia - {os.path.basename(file_path)}")
            
            view_window.after(0, lambda: view_window.state("zoomed"))  # Garantir foco na nova janela
            view_window.bind("<Escape>", lambda e: view_window.destroy())  # Fechar com ESC
            
            view_window.configure(fg_color="#2b2b2b")  # Fundo escuro Excel

            # Frame com Scroll
            scroll_frame = customtkinter.CTkScrollableFrame(
                view_window,
                fg_color="#1e1e1e",  # Fundo escuro
                border_width=1,
                border_color="#4F81BD"  # Borda azul Excel
            )
            scroll_frame.pack(fill="both", expand=True, padx=20, pady=20)

            # 2. CONFIGURAR COLUNAS
            # Lemos a primeira linha para saber quantas colunas existem
            headers = [cell.value for cell in sheet[1]] # type: ignore
            for i in range(len(headers)):
                scroll_frame.grid_columnconfigure(i, weight=1, minsize=150)

            # 3. RENDERIZAR CABEÇALHO com cores Excel
            for i, header in enumerate(headers):
                h_label = customtkinter.CTkLabel(
                    scroll_frame, 
                    text=str(header or "").upper(), 
                    font=self.fonts.TABLE_HEADER,
                    fg_color="#4F81BD",  # Azul Excel header
                    text_color="white",
                    corner_radius=5,
                    height=35
                )
                h_label.grid(row=0, column=i, padx=2, pady=10, sticky="nsew")

            # 4. RENDERIZAR DADOS
            # iter_rows começando da linha 2 (para pular o cabeçalho)
            for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=1): # type: ignore
                for col_idx, cell_value in enumerate(row):
                    val_str = str(cell_value or "")
                    
                    # Cores baseadas no conteúdo (Sua estratégia de controle)
                    txt_color = "white"
                    if "Pago" in val_str or "Sim" in val_str:
                        txt_color = "#2ecc71" # Verde para sucesso
                    elif "Pendente" in val_str or "Não" in val_str:
                        txt_color = "#e67e22" # Laranja para atenção
                    
                    b_label = customtkinter.CTkLabel(
                        scroll_frame, 
                        text=val_str,
                        font=self.fonts.TABLE_CELL,
                        text_color=txt_color,
                        fg_color="#3a3a3a",  # Fundo escuro para células
                        height=35
                    )
                    b_label.grid(row=row_idx, column=col_idx, padx=1, pady=1, sticky="nsew")

        except Exception as e:
            self.result_textbox.delete("0.0", "end")
            self.result_textbox.insert("0.0", f"Erro ao abrir {os.path.basename(file_path)}: {str(e)}")
    
    def _parse_currency(self, value):
        """
        Converte um valor para float, tratando diferentes formatos de moeda.
        
        Args:
            value: Valor a ser convertido (string ou numérico)
            
        Returns:
            float: Valor convertido ou 0.0 se inválido
        """
        if value is None:
            return 0.0
        try:
            # Remove símbolos de moeda e espaços
            str_value = str(value).replace('R$', '').replace('$', '').replace(' ', '').replace(',', '.')
            return float(str_value)
        except (ValueError, TypeError):
            return 0.0